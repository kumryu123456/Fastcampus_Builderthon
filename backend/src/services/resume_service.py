"""
Resume service for file upload, text extraction, and analysis.

Constitution Compliance:
- Principle II: API Resilience - Retry logic for external calls
- Principle III: User Data Privacy - Secure file storage, PII scrubbing
- Principle V: Code Quality - Structured logging throughout

T024: Resume upload and text extraction
T025: Resume analysis orchestration
T026: Caching to avoid duplicate Gemini API calls
"""

import os
import uuid
import time
from pathlib import Path
from typing import Optional, Dict, Any, BinaryIO
from datetime import datetime

import PyPDF2
import docx
from fastapi import UploadFile
from sqlalchemy.orm import Session

from src.config import settings
from src.models.resume import Resume
from src.models.user import User
from src.services.gemini_client import GeminiClient
from src.utils.logging_config import get_logger
from src.utils.privacy import scrub_all_pii

logger = get_logger(__name__)


class ResumeService:
    """
    Service for handling resume operations.

    T024: File upload and text extraction
    T025: Analysis orchestration
    T026: Caching logic
    """

    # Allowed file types
    ALLOWED_MIME_TYPES = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",  # Legacy .doc
        "application/octet-stream": None,  # Sometimes returned for .docx
    }
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    def __init__(self, db: Session):
        """
        Initialize resume service.

        Args:
            db: Database session
        """
        self.db = db
        self.gemini_client = GeminiClient()
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def upload_and_analyze_resume(
        self,
        file: UploadFile,
        user_id: int,
    ) -> Resume:
        """
        Complete workflow: Upload file → Extract text → Analyze → Cache results.

        T024: File upload and validation
        T025: Text extraction and analysis
        T026: Check cache before calling Gemini

        Args:
            file: Uploaded file
            user_id: User ID

        Returns:
            Resume object with analysis results

        Raises:
            ValueError: If file is invalid
            Exception: If analysis fails
        """
        start_time = time.time()

        logger.info(
            "resume_upload_started",
            operation="upload_and_analyze",
            user_id=f"user-{user_id}",
            filename=scrub_all_pii(file.filename),
            content_type=file.content_type,
        )

        try:
            # Step 1: Validate file (T024)
            self._validate_file(file)

            # Step 2: Read file content
            file_content = await file.read()
            file_size = len(file_content)
            file_hash = Resume.compute_file_hash(file_content)

            logger.info(
                "file_read_completed",
                operation="upload_and_analyze",
                user_id=f"user-{user_id}",
                file_size=file_size,
                file_hash=file_hash[:16],  # First 16 chars
            )

            # Step 3: Check cache (T026) - Avoid duplicate Gemini calls
            cached_resume = self._get_cached_resume(user_id, file_hash)
            if cached_resume:
                logger.info(
                    "resume_analysis_cache_hit",
                    operation="upload_and_analyze",
                    user_id=f"user-{user_id}",
                    cached_resume_id=cached_resume.id,
                )
                return cached_resume

            # Step 4: Save file to disk (Constitution III: UUID filename)
            file_path = self._save_file(file.filename, file_content)

            # Step 5: Create database record
            resume = Resume(
                user_id=user_id,
                original_filename=file.filename,
                file_path=str(file_path),
                file_size=file_size,
                mime_type=file.content_type,
                file_hash=file_hash,
                status="processing",
            )
            self.db.add(resume)
            self.db.commit()
            self.db.refresh(resume)

            logger.info(
                "resume_record_created",
                operation="upload_and_analyze",
                user_id=f"user-{user_id}",
                resume_id=resume.id,
            )

            # Step 6: Extract text from file (T024)
            try:
                extracted_text = self._extract_text(file_path, file.content_type)
                resume.extracted_text = extracted_text
                self.db.commit()

                logger.info(
                    "text_extraction_completed",
                    operation="upload_and_analyze",
                    resume_id=resume.id,
                    text_length=len(extracted_text),
                )
            except Exception as e:
                resume.status = "failed"
                resume.error_message = f"Text extraction failed: {str(e)}"
                self.db.commit()
                raise

            # Step 7: Analyze with Gemini (T025)
            try:
                analysis_result = self.gemini_client.analyze_resume_text(
                    extracted_text,
                    user_id=user_id,
                )

                resume.analysis_result = analysis_result
                resume.status = "analyzed"
                resume.analyzed_at = datetime.utcnow()
                self.db.commit()

                logger.info(
                    "resume_analysis_success",
                    operation="upload_and_analyze",
                    resume_id=resume.id,
                    user_id=f"user-{user_id}",
                )

            except Exception as e:
                resume.status = "failed"
                resume.error_message = f"Analysis failed: {str(e)}"
                self.db.commit()
                raise

            duration_ms = int((time.time() - start_time) * 1000)
            logger.info(
                "resume_upload_completed",
                operation="upload_and_analyze",
                resume_id=resume.id,
                user_id=f"user-{user_id}",
                duration_ms=duration_ms,
                status=resume.status,
            )

            return resume

        except Exception as e:
            duration_ms = int((time.time() - start_time) * 1000)
            logger.error(
                "resume_upload_failed",
                operation="upload_and_analyze",
                user_id=f"user-{user_id}",
                duration_ms=duration_ms,
                error=str(e),
                error_type=type(e).__name__,
                exc_info=True,
            )
            raise

    def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file.

        T024: File validation (type, size)

        Args:
            file: Uploaded file

        Raises:
            ValueError: If file is invalid
        """
        # Get file extension
        file_ext = Path(file.filename).suffix.lower() if file.filename else ""

        # Check file type (MIME type OR extension)
        is_valid_mime = file.content_type in self.ALLOWED_MIME_TYPES
        is_valid_ext = file_ext in self.ALLOWED_EXTENSIONS

        if not (is_valid_mime or is_valid_ext):
            raise ValueError(
                f"Invalid file type: {file.content_type} ({file_ext}). "
                f"Allowed types: PDF, DOCX"
            )

        # File size will be checked after reading
        # (UploadFile doesn't provide size before reading)

        logger.debug(
            "file_validation_passed",
            operation="validate_file",
            content_type=file.content_type,
            file_extension=file_ext,
        )

    def _save_file(self, original_filename: str, content: bytes) -> Path:
        """
        Save file to disk with UUID filename.

        Constitution III: Secure storage - UUID filenames prevent information leakage

        Args:
            original_filename: Original filename
            content: File content

        Returns:
            Path to saved file

        Raises:
            ValueError: If file is too large
        """
        # Check file size (Constitution requirement: <5MB)
        file_size_mb = len(content) / (1024 * 1024)
        if file_size_mb > settings.max_upload_size_mb:
            raise ValueError(
                f"File too large: {file_size_mb:.2f}MB. "
                f"Maximum allowed: {settings.max_upload_size_mb}MB"
            )

        # Generate UUID filename with original extension
        file_ext = Path(original_filename).suffix
        uuid_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / uuid_filename

        # Save file
        with open(file_path, "wb") as f:
            f.write(content)

        logger.info(
            "file_saved",
            operation="save_file",
            file_path=str(file_path),
            file_size_mb=file_size_mb,
        )

        return file_path

    def _extract_text(self, file_path: Path, mime_type: str) -> str:
        """
        Extract text from PDF or DOCX file.

        T024: Text extraction using PyPDF2 and python-docx

        Args:
            file_path: Path to file
            mime_type: MIME type

        Returns:
            Extracted text

        Raises:
            ValueError: If extraction fails
        """
        try:
            # Determine file type by extension if MIME type is ambiguous
            file_ext = file_path.suffix.lower()

            if mime_type == "application/pdf" or file_ext == ".pdf":
                return self._extract_text_from_pdf(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_ext in [".docx", ".doc"]:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {mime_type} ({file_ext})")
        except Exception as e:
            logger.error(
                "text_extraction_failed",
                operation="extract_text",
                file_path=str(file_path),
                mime_type=mime_type,
                error=str(e),
                exc_info=True,
            )
            raise ValueError(f"Failed to extract text: {str(e)}")

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using PyPDF2.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text_parts = []

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                logger.debug(
                    "pdf_page_extracted",
                    operation="extract_pdf",
                    page_num=page_num + 1,
                    text_length=len(page_text) if page_text else 0,
                )

            full_text = "\n".join(text_parts)

            logger.info(
                "pdf_extraction_completed",
                operation="extract_pdf",
                total_pages=len(pdf_reader.pages),
                total_text_length=len(full_text),
            )

            return full_text

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        doc = docx.Document(file_path)
        text_parts = []

        for para_num, para in enumerate(doc.paragraphs):
            if para.text:
                text_parts.append(para.text)

            logger.debug(
                "docx_paragraph_extracted",
                operation="extract_docx",
                para_num=para_num + 1,
                text_length=len(para.text) if para.text else 0,
            )

        full_text = "\n".join(text_parts)

        logger.info(
            "docx_extraction_completed",
            operation="extract_docx",
            total_paragraphs=len(doc.paragraphs),
            total_text_length=len(full_text),
        )

        return full_text

    def _get_cached_resume(self, user_id: int, file_hash: str) -> Optional[Resume]:
        """
        Check if resume with same hash already exists for user.

        T026: Caching logic to avoid duplicate Gemini API calls
        Constitution II: Optimization reduces API costs and latency

        Args:
            user_id: User ID
            file_hash: File hash

        Returns:
            Cached resume if found (with valid ats_score), None otherwise
        """
        cached = (
            self.db.query(Resume)
            .filter(
                Resume.user_id == user_id,
                Resume.file_hash == file_hash,
                Resume.status == "analyzed",
            )
            .first()
        )

        if cached:
            # Check if cached result has ATS score (new feature)
            # If not, invalidate cache to trigger re-analysis
            if cached.analysis_result and "ats_score" not in cached.analysis_result:
                logger.info(
                    "resume_cache_invalidated_no_ats",
                    operation="get_cached_resume",
                    user_id=f"user-{user_id}",
                    cached_resume_id=cached.id,
                    file_hash=file_hash[:16],
                    reason="missing_ats_score",
                )
                return None

            logger.info(
                "resume_cache_hit",
                operation="get_cached_resume",
                user_id=f"user-{user_id}",
                cached_resume_id=cached.id,
                file_hash=file_hash[:16],
            )

        return cached

    def get_resume_by_id(self, resume_id: int, user_id: int) -> Optional[Resume]:
        """
        Get resume by ID for specific user.

        Args:
            resume_id: Resume ID
            user_id: User ID

        Returns:
            Resume object or None
        """
        return (
            self.db.query(Resume)
            .filter(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
            .first()
        )
